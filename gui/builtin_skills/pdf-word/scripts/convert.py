"""PDF to Word converter using PyMuPDF + python-docx."""
import argparse
import os
import re
import sys
import tempfile


# XML 1.0 allows: #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]
# Remove all characters outside this set to prevent lxml ValueError.
_ILLEGAL_XML_CHARS_RE = re.compile(
    r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f'
    r'\ud800-\udfff\ufdd0-\ufdef\ufffe\uffff]'
)


def _clean_xml_text(text: str) -> str:
    """Remove characters that are not valid in XML 1.0."""
    if not text:
        return text
    return _ILLEGAL_XML_CHARS_RE.sub('', text)

def convert_pdf_to_docx(input_path: str, output_path: str) -> str:
    """Convert a PDF file to DOCX preserving text and basic layout."""
    try:
        import fitz  # pymupdf
    except ImportError:
        print("ERROR: pymupdf not installed. Run: pip install pymupdf", file=sys.stderr)
        sys.exit(1)
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    if not os.path.isfile(input_path):
        print(f"ERROR: Input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    # Determine output path
    if os.path.isdir(output_path):
        base = os.path.splitext(os.path.basename(input_path))[0]
        output_path = os.path.join(output_path, f"{base}.docx")
    elif not output_path.lower().endswith('.docx'):
        output_path += '.docx'

    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)

    doc = Document()
    pdf = fitz.open(input_path)

    if pdf.is_encrypted:
        print("ERROR: PDF is encrypted/password-protected. Please provide an unprotected PDF.", file=sys.stderr)
        pdf.close()
        sys.exit(1)

    page_count = len(pdf)
    print(f"Starting conversion: {page_count} page(s)", flush=True)

    for page_num in range(page_count):
        print(f"Processing page {page_num + 1}/{page_count}", flush=True)
        page = pdf[page_num]
        blocks = page.get_text("dict")["blocks"]
        page_has_images = False

        for block in blocks:
            if block["type"] == 0:  # text block
                for line in block.get("lines", []):
                    text_parts = []
                    for span in line.get("spans", []):
                        text_parts.append(span["text"])
                    line_text = "".join(text_parts).strip()
                    # Clean illegal XML characters that may exist in PDF text
                    line_text = _clean_xml_text(line_text)
                    if line_text:
                        para = doc.add_paragraph()
                        # Detect heading by font size
                        max_size = max((span.get("size", 11) for span in line.get("spans", [])), default=11)
                        if max_size >= 18:
                            para.style = 'Heading 1'
                        elif max_size >= 14:
                            para.style = 'Heading 2'
                        else:
                            para.style = 'Normal'
                        
                        for span in line.get("spans", []):
                            run = para.add_run(_clean_xml_text(span["text"]))
                            run.font.size = Pt(span.get("size", 11))
                            flags = span.get("flags", 0)
                            if flags & 16:  # bit 4 = bold
                                run.bold = True
                            if flags & 2:   # bit 1 = italic
                                run.italic = True

            elif block["type"] == 1:  # image block marker
                page_has_images = True

        # Extract and embed page images once (after all blocks processed)
        if page_has_images:
            for img_info in page.get_images(full=True):
                img_path = None
                try:
                    xref = img_info[0]
                    base_image = pdf.extract_image(xref)
                    if not base_image or not base_image.get("image"):
                        continue
                    # Use a unique system temp file instead of an output-folder
                    # name based on page/xref: concurrent conversions can share
                    # those values and previously raced to overwrite/delete it.
                    img_ext = re.sub(r"[^A-Za-z0-9]", "", base_image.get("ext", "png")) or "png"
                    with tempfile.NamedTemporaryFile(prefix="maclaw-pdf-word-", suffix=f".{img_ext}", delete=False) as temp_file:
                        img_path = temp_file.name
                        temp_file.write(base_image["image"])
                    doc.add_picture(img_path, width=Inches(5))
                except Exception:
                    # A broken embedded image should not discard extracted text
                    # or fail the entire document conversion.
                    pass
                finally:
                    if img_path:
                        try:
                            os.remove(img_path)
                        except OSError:
                            pass

        # Add page break between pages (except last)
        if page_num < page_count - 1:
            doc.add_page_break()

    pdf.close()
    doc.save(output_path)
    print(f"Conversion complete: {output_path}", flush=True)
    print(f"Pages: {page_count}", flush=True)
    print(f"Output: {os.path.abspath(output_path)}", flush=True)
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Convert PDF to Word (DOCX)")
    parser.add_argument("--input", required=True, help="Input PDF file path")
    parser.add_argument("--output", default=".", help="Output DOCX file path or directory")
    args = parser.parse_args()
    convert_pdf_to_docx(args.input, args.output)


if __name__ == "__main__":
    main()
