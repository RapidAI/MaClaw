"""Document PII redaction using regex patterns for Chinese and international PII."""
import argparse
import os
import re
import sys

# --- PII Patterns ---

PATTERNS = [
    # 中国身份证号（18位）
    (r'\b[1-9]\d{5}(?:19|20)\d{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12]\d|3[01])\d{3}[\dXx]\b',
     '[身份证号已脱敏]'),
    # 中国手机号
    (r'\b1[3-9]\d{9}\b', '[手机号已脱敏]'),
    # 固定电话
    (r'\b0\d{2,3}[-\s]?\d{7,8}\b', '[电话已脱敏]'),
    # 邮箱
    (r'\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b', '[邮箱已脱敏]'),
    # 银行卡号（严格：以常见BIN开头的16-19位数字，支持空格/横线分隔）
    (r'\b(?:62|4\d|5[1-5]|3[47])\d{2}[\s\-]?\d{4}[\s\-]?\d{4}[\s\-]?\d{3,6}\b',
     '[银行卡号已脱敏]'),
    # IP 地址（排除版本号如 1.2.3）
    (r'\b(?:(?:25[0-5]|2[0-4]\d|1\d{2}|[1-9]?\d)\.){3}(?:25[0-5]|2[0-4]\d|1\d{2}|[1-9]?\d)\b',
     '[IP地址已脱敏]'),
    # 中国地址（要求"省/市/自治区"开头 + 至少市/区/县/路级别细化）
    (r'(?:[\u4e00-\u9fa5]{2,}(?:省|自治区|市))[\u4e00-\u9fa5]{1,}(?:市|区|县|镇|乡)[\u4e00-\u9fa5\d\-]*(?:路|街|巷|号|栋|单元|室|楼|层|弄|村)[\u4e00-\u9fa5\d\-]*',
     '[地址已脱敏]'),
]

COMPILED_PATTERNS = [(re.compile(p), repl) for p, repl in PATTERNS]


def redact_text(text: str) -> "tuple[str, int]":
    """Apply all PII patterns to text, return (redacted_text, count)."""
    total = 0
    for pattern, replacement in COMPILED_PATTERNS:
        text, n = pattern.subn(replacement, text)
        total += n
    return text, total


def redact_txt_file(input_path: str, output_path: str) -> int:
    """Redact a plain text or markdown file."""
    # Try UTF-8 first, fall back to GBK (common in Chinese enterprise docs)
    content = None
    for enc in ('utf-8', 'utf-8-sig', 'gbk', 'gb2312', 'latin-1'):
        try:
            with open(input_path, 'r', encoding=enc) as f:
                content = f.read()
            break
        except (UnicodeDecodeError, ValueError):
            continue
    if content is None:
        print(f"ERROR: Cannot decode {input_path} with any supported encoding", file=sys.stderr)
        sys.exit(1)
    
    redacted, count = redact_text(content)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(redacted)
    
    return count


def redact_docx_file(input_path: str, output_path: str) -> int:
    """Redact a DOCX file preserving formatting where possible.
    
    Strategy: try run-level replacement first (preserves formatting).
    If a paragraph contains PII that spans across runs (not caught by run-level),
    fall back to paragraph-level replacement (loses formatting for that paragraph).
    """
    try:
        from docx import Document
    except ImportError:
        print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
        sys.exit(1)
    
    doc = Document(input_path)
    total = 0

    def redact_paragraph(para) -> int:
        """Redact a single paragraph, preferring run-level, falling back to paragraph-level."""
        count = 0
        # Pass 1: run-level (preserves formatting)
        for run in para.runs:
            if run.text:
                new_text, n = redact_text(run.text)
                if n > 0:
                    run.text = new_text
                    count += n
        
        # Pass 2: check if full paragraph still has PII (cross-run patterns)
        full_text = para.text
        if full_text:
            redacted_full, remaining = redact_text(full_text)
            if remaining > 0:
                # PII spans across runs — fall back to paragraph-level replacement
                # Clear all runs and set text on first run (loses formatting for this paragraph)
                for i, run in enumerate(para.runs):
                    if i == 0:
                        run.text = redacted_full
                    else:
                        run.text = ""
                count += remaining
        
        return count
    
    for para in doc.paragraphs:
        total += redact_paragraph(para)
    
    # Also process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    total += redact_paragraph(para)
    
    doc.save(output_path)
    return total


def main():
    parser = argparse.ArgumentParser(description="Redact PII from documents")
    parser.add_argument("--input", required=True, help="Input file (TXT/MD/DOCX)")
    parser.add_argument("--output", default=".", help="Output file path or directory")
    args = parser.parse_args()

    input_path = args.input
    if not os.path.isfile(input_path):
        print(f"ERROR: Input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    ext = os.path.splitext(input_path)[1].lower()
    base = os.path.splitext(os.path.basename(input_path))[0]

    # Determine output path
    output_path = args.output
    if os.path.isdir(output_path):
        output_path = os.path.join(output_path, f"{base}_redacted{ext}")
    
    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)

    if ext == '.docx':
        count = redact_docx_file(input_path, output_path)
    elif ext in ('.txt', '.md', '.markdown', '.csv', '.log'):
        count = redact_txt_file(input_path, output_path)
    else:
        print(f"ERROR: Unsupported file format: {ext}. Supported: .txt .md .docx .csv .log", file=sys.stderr)
        sys.exit(1)

    print(f"脱敏完成: {output_path}")
    print(f"   替换数量: {count} 处敏感信息")
    print(f"   输出文件: {os.path.abspath(output_path)}")


if __name__ == "__main__":
    main()
