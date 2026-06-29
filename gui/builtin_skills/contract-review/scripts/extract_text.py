"""Extract text from contract files (PDF/DOCX/TXT) for AI review."""
import argparse
import os
import sys


def extract_from_pdf(path: str) -> str:
    """Extract text from PDF using PyMuPDF."""
    try:
        import fitz
    except ImportError:
        print("ERROR: pymupdf not installed. Run: pip install pymupdf", file=sys.stderr)
        sys.exit(1)
    
    doc = fitz.open(path)
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()
    return "\n".join(text_parts)


def extract_from_docx(path: str) -> str:
    """Extract text from DOCX."""
    try:
        from docx import Document
    except ImportError:
        print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
        sys.exit(1)
    
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def extract_from_txt(path: str) -> str:
    """Read plain text file with encoding auto-detection."""
    for enc in ('utf-8', 'utf-8-sig', 'gbk', 'gb2312', 'latin-1'):
        try:
            with open(path, 'r', encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, ValueError):
            continue
    # Last resort: read as binary and decode lossily
    with open(path, 'rb') as f:
        return f.read().decode('utf-8', errors='replace')


def main():
    parser = argparse.ArgumentParser(description="Extract text from contract files")
    parser.add_argument("--input", required=True, help="Input file (PDF/DOCX/TXT)")
    parser.add_argument("--output", default=".", help="Output directory for extracted text")
    args = parser.parse_args()

    input_path = args.input
    if not os.path.isfile(input_path):
        print(f"ERROR: File not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    ext = os.path.splitext(input_path)[1].lower()
    if ext == '.pdf':
        text = extract_from_pdf(input_path)
    elif ext == '.docx':
        text = extract_from_docx(input_path)
    elif ext in ('.txt', '.md'):
        text = extract_from_txt(input_path)
    else:
        print(f"ERROR: Unsupported format: {ext}. Use PDF, DOCX, or TXT.", file=sys.stderr)
        sys.exit(1)

    os.makedirs(args.output, exist_ok=True)
    base = os.path.splitext(os.path.basename(input_path))[0]
    output_path = os.path.join(args.output, f"{base}_text.txt")
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)

    print(f"✅ 文本提取完成")
    print(f"   字符数: {len(text)}")
    print(f"   output_path: {output_path}")


if __name__ == "__main__":
    main()
